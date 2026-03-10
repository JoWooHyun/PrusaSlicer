from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 기본 스타일 설정
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph(line)
        p.style = doc.styles['No Spacing']
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

def add_table_with_header(doc, data, style_name='Light Shading Accent 1'):
    table = doc.add_table(rows=len(data), cols=len(data[0]), style=style_name)
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = str(cell_text)
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    return table

def add_bold_then_normal(doc, bold_text, normal_text):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    p.add_run(normal_text)
    return p

# ============================================================
# 표지
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('하이브리드 멀티레진 3D 프린터\n커스텀 슬라이서 개발 계획서')
run.font.size = Pt(26)
run.bold = True

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Top-Down 방식 | FDM 도포 + DLP 마스크 경화 | 멀티레진')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('기반 소스: PrusaSlicer v2.9.4 (AGPLv3)\n참고: hybrid_resin_slicer_guideline.docx\n작성일: 2026-03-10')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ============================================================
# 목차
# ============================================================
doc.add_heading('목차', level=1)
toc_items = [
    '1. 프로젝트 개요',
    '   1.1 프린터 방식',
    '   1.2 레이어 공정 (5단계)',
    '   1.3 멀티 레진 개념',
    '   1.4 사용 워크플로우',
    '   1.5 하드웨어 환경',
    '2. 슬라이서 설계 방향',
    '   2.1 핵심 관점: 레진 슬라이서가 몸통',
    '   2.2 슬라이싱 파이프라인 (하이브리드)',
    '   2.3 출력 데이터 구조',
    '   2.4 설계 원칙',
    '3. PrusaSlicer에서 활용할 부분',
    '   3.1 SLA 쪽에서 가져올 것 (몸통)',
    '   3.2 FDM 쪽에서 가져올 것 (추가)',
    '   3.3 제거할 것',
    '   3.4 참고 오픈소스 프로젝트',
    '4. 개발 Phase',
    '   Phase 0: 빌드 환경 구축',
    '   Phase 1: 불필요 기능 제거 + 브랜딩',
    '   Phase 2: 멀티 레진 Tool 매핑',
    '   Phase 3: 마스크 이미지(PNG) 생성',
    '   Phase 4: 도포 경로 GCode 생성',
    '   Phase 5: 출력 포맷 패키징 (ZIP)',
    '   Phase 6: vgui 연동 + 실 프린트 테스트',
    '5. 주요 수정 파일 목록',
    '6. vgui (프린터 측) 수정 사항',
    '7. 핵심 원칙 + 일정',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)

doc.add_page_break()

# ============================================================
# 1. 프로젝트 개요
# ============================================================
doc.add_heading('1. 프로젝트 개요', level=1)

doc.add_heading('1.1 프린터 방식', level=2)
doc.add_paragraph(
    'Top-Down 방식의 하이브리드 멀티레진 3D 프린터를 위한 커스텀 슬라이서를 개발한다. '
    'PrusaSlicer(v2.9.4) 소스코드를 기반으로 하되, 레진 도포 시스템에 맞게 수정한다.'
)
doc.add_paragraph('프린터는 세 가지 기술을 결합한다:')
for b in [
    'FDM 방식의 도포 경로 (XYZ 이동으로 레진 공급)',
    '블레이드 평탄화 (도포된 레진을 균일하게 펼침)',
    'DLP 마스크 경화 (LED 프로젝터로 선택적 경화)',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph()
add_bold_then_normal(doc,
    '핵심 포인트: ',
    '도포 경로(GCode)는 재료 공급만 담당하고, 최종 형상의 정밀도는 마스크 이미지(PNG)가 결정한다.')

doc.add_heading('1.2 레이어 공정 (5단계)', level=2)
doc.add_paragraph('각 레이어는 다음 5단계로 처리된다:')

add_table_with_header(doc, [
    ('단계', '동작', '출력 데이터'),
    ('1. 레진 도포 (T0)', '모델레진 영역에 T0 노즐로 도포', 'GCode (G1 X.. Y.. E..)'),
    ('2. 레진 도포 (T1)', '템프레진 영역에 T1 노즐로 도포', 'GCode (G1 X.. Y.. E..)'),
    ('3. 블레이드 평탄화', '블레이드로 레진 표면을 고르게 함', 'GCode (블레이드 스윕)'),
    ('4. 마스크 노광', '프로젝터가 마스크 이미지를 조사하여 경화', 'PNG (1장, 전체 합침)'),
    ('5. Z축 이동', '다음 레이어 높이로 이동', 'GCode (G1 Z..)'),
])

doc.add_paragraph()
add_bold_then_normal(doc,
    '가정 (이상적 시나리오): ',
    '두 레진의 경화 시간 동일, 블레이드 평탄화 시 레진 혼합 없음, 도포 영역 정확히 분리됨. '
    '물리적 문제는 하드웨어/재료 쪽에서 해결할 영역이며, 소프트웨어는 이상적으로 동작한다고 전제한다.')

doc.add_heading('1.3 멀티 레진 개념', level=2)
doc.add_paragraph(
    '여러 STL 파일을 동일 좌표계에 불러온다. '
    'STL 파일들은 서로 겹칠 수 있다. (예: 잇몸 모델 구멍 안에 치아 모델이 들어감)'
)

doc.add_paragraph('예시 (치과 모델):')
add_table_with_header(doc, [
    ('STL 파일', 'Tool', '레진', '역할'),
    ('잇몸_모델.stl', 'T0', '모델레진', '잇몸 구조체'),
    ('치아_21.stl', 'T1', '템프레진', '치아 복원체'),
    ('치아_22.stl', 'T1', '템프레진', '치아 복원체'),
])

doc.add_paragraph()
doc.add_paragraph(
    '모델들은 외부 CAD/치과 소프트웨어(exocad 등)에서 좌표가 이미 정해진 상태로 준비된다. '
    '슬라이서는 STL을 수정하지 않고 그대로 슬라이스한다. '
    '사용자가 수동으로 배치할 필요 없이, 외부에서 잡힌 좌표 그대로 불러오면 된다.')

doc.add_paragraph()
doc.add_paragraph('레진 혼합 정책:')
for b in [
    '레진 혼합은 허용된다',
    '경화되지 않은 혼합 레진은 재사용하지 않고 폐기한다',
    '시스템의 목표는 완벽한 재료 분리가 아니라 멀티 레진 출력 기능 구현이다',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('1.4 사용 워크플로우', level=2)
add_code_block(doc, [
    '[외부 CAD]              [PC]                    [USB]           [프린터]',
    'exocad 등에서      커스텀 슬라이서에서       USB에 저장       vgui가 파일을',
    'STL 파일 준비  ->  로드 + Tool 지정    ->  ZIP 파일   ->  읽어서 프린팅',
    '(좌표 확정)        슬라이스                                실행',
])

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('슬라이서는 PC 전용 데스크탑 앱이며, 프린터와 직접 통신하지 않음 (오프라인 USB 전달)')
run.bold = True

doc.add_heading('1.5 하드웨어 환경', level=2)
add_table_with_header(doc, [
    ('구분', '사양'),
    ('보드', 'BigTreeTech Manta M8P 2.0'),
    ('컴퓨트', 'Raspberry Pi CM4'),
    ('펌웨어', 'Klipper + Moonraker'),
    ('프린터 GUI', 'vgui (PySide6, 자체 개발)'),
    ('프로젝터', 'DLP LED 프로젝터 (마스크 경화용)'),
    ('도포 노즐', '멀티 노즐 (T0, T1, ...)'),
    ('블레이드', '레진 평탄화용'),
    ('현재 슬라이서', '치투슬라이서 (교체 예정)'),
])

doc.add_page_break()

# ============================================================
# 2. 슬라이서 설계 방향
# ============================================================
doc.add_heading('2. 슬라이서 설계 방향', level=1)

doc.add_heading('2.1 핵심 관점: 레진 슬라이서가 몸통', level=2)

p = doc.add_paragraph()
run = p.add_run('중요: ')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run('이 슬라이서의 몸통은 FDM이 아니라 레진(SLA) 슬라이서이다.')

doc.add_paragraph()
add_code_block(doc, [
    '잘못된 접근:',
    '  FDM 슬라이서가 몸통 + SLA 기능 추가',
    '',
    '올바른 접근:',
    '  레진 슬라이서가 몸통 + FDM 도포 기능 추가',
    '',
    '이유:',
    '  - 최종 출력물의 품질을 결정하는 건 마스크 이미지(PNG)이다',
    '  - 도포 경로(GCode)는 "어디에 레진을 뿌릴지" 수준의 대략적 경로이다',
    '  - 서포트, 미리보기 등 서브기능도 레진 슬라이서 기반이어야 한다',
])

doc.add_paragraph()
doc.add_paragraph('따라서 구조는:')
add_code_block(doc, [
    '┌─────────────────────────────────────────┐',
    '│  레진 슬라이서 (몸통)                     │',
    '│  ├─ 3D 뷰어 + STL 배치                   │',
    '│  ├─ 레이어 슬라이싱 (Z 절단)              │',
    '│  ├─ 마스크 이미지 생성 (PNG)              │ ← 핵심',
    '│  ├─ 서포트 생성                           │',
    '│  ├─ 미리보기                              │',
    '│  │                                        │',
    '│  ├─ [추가] FDM식 도포 경로 생성 (GCode)   │',
    '│  ├─ [추가] 멀티 레진 Tool 지정            │',
    '│  ├─ [추가] 블레이드 평탄화 GCode          │',
    '│  └─ [추가] ZIP 패키징                     │',
    '└─────────────────────────────────────────┘',
])

doc.add_heading('2.2 슬라이싱 파이프라인 (하이브리드)', level=2)

add_code_block(doc, [
    'STL 여러 개 로드 (외부 CAD에서 좌표 확정 상태)',
    '    |',
    '    v',
    'STL별 레진(Tool) 지정 (T0=모델레진, T1=템프레진)',
    '    |',
    '    v',
    'Z 슬라이싱 (레이어 높이별로 절단 -> ExPolygons)',
    '    |',
    '    +---> 마스크 이미지 생성 (전체 합쳐서 1장 PNG)',
    '    |',
    '    +---> Tool별 도포 영역 -> 도포 경로 GCode 생성',
    '    |',
    '    +---> 블레이드/Z이동 GCode 삽입',
    '    |',
    '    v',
    'ZIP 패키징 (run.gcode + layers/*.png + manifest.json)',
])

doc.add_heading('2.3 출력 데이터 구조', level=2)
doc.add_paragraph('슬라이서는 두 가지 데이터를 생성한다:')
doc.add_paragraph('1. GCode: 도포 경로 + 블레이드 + Z이동 (Klipper 호환)')
doc.add_paragraph('2. PNG 마스크 이미지: 레이어별 경화 패턴 (재료 구분 없이 합쳐진 1장)')

doc.add_paragraph()
doc.add_paragraph('출력 파일 구조:')
add_code_block(doc, [
    'print_job.zip',
    '├── run.gcode              <- 전체 도포/블레이드/Z이동 명령',
    '├── layers/',
    '│   ├── 0000.png           <- 레이어 0 마스크 (흰=경화, 검=미경화)',
    '│   ├── 0001.png           <- 레이어 1 마스크',
    '│   ├── 0002.png',
    '│   └── ...',
    '├── preview.png            <- 미리보기 썸네일',
    '└── manifest.json          <- 메타데이터',
])

doc.add_paragraph()
doc.add_paragraph('마스크 이미지 설명:')
add_code_block(doc, [
    '┌──────────────────────────┐',
    '│ ████████████████████████ │  흰색 = LED 켜짐 = 경화',
    '│ ████  ░░░░  ░░░░  ████  │  검은색 = LED 꺼짐 = 미경화',
    '│ ████  ░░░░  ░░░░  ████  │',
    '│ ████████████████████████ │  ████ = 잇몸 영역 (T0 모델레진)',
    '│                          │  ░░░░ = 치아 영역 (T1 템프레진)',
    '└──────────────────────────┘',
    '',
    '마스크는 재료 구분 없이 "굳힐 곳 = 흰색" 한 장이면 끝.',
    '어디에 어떤 레진을 뿌릴지는 GCode가 담당한다.',
])

doc.add_heading('2.4 설계 원칙', level=2)
for b in [
    '슬라이서는 단순하고 안정적인 구조를 유지해야 한다',
    'Boolean 연산 등 복잡한 메시 연산은 피한다',
    'STL 모델 준비는 외부 소프트웨어(exocad 등)에서 담당한다',
    '슬라이서는 STL을 수정하지 않고 그대로 슬라이스한다',
    '기존 기능 위에 추가하는 방식이지, 처음부터 새로 만들지 않는다',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ============================================================
# 3. PrusaSlicer에서 활용할 부분
# ============================================================
doc.add_heading('3. PrusaSlicer에서 활용할 부분', level=1)

doc.add_heading('3.1 SLA 쪽에서 가져올 것 (몸통)', level=2)
add_table_with_header(doc, [
    ('기능', '소스 위치', '용도'),
    ('3D 뷰어 + 모델 배치', 'src/slic3r/GUI/', '기존 UI 그대로 사용'),
    ('Z 슬라이싱 -> ExPolygon', 'PrintObjectSlice.cpp', '레이어 단면 생성'),
    ('ExPolygon -> PNG 래스터화', 'src/libslic3r/SLA/RasterBase.*', '마스크 이미지 생성'),
    ('SLA 서포트', 'src/libslic3r/SLA/', '서포트 구조 (레진 방식)'),
    ('SL1 ZIP 출력', 'src/libslic3r/Format/SL1.*', 'ZIP 패키징 참고'),
    ('레이어 높이/노출 시간 설정', 'PrintConfig.hpp (SLA 파라미터)', '경화 설정'),
    ('Zipper 클래스', 'src/libslic3r/Zipper.*', 'ZIP 파일 생성'),
])

doc.add_heading('3.2 FDM 쪽에서 가져올 것 (추가)', level=2)
add_table_with_header(doc, [
    ('기능', '소스 위치', '용도'),
    ('멀티 익스트루더 -> 멀티 레진 Tool', 'PrintConfig.hpp, PrintObject.cpp', 'STL별 Tool 지정'),
    ('인필 패턴 -> 도포 경로', 'src/libslic3r/Fill/', '레진 도포 패턴으로 활용'),
    ('GCode 생성', 'src/libslic3r/GCode/*', 'Klipper용 도포 GCode'),
    ('GCodeWriter (gcfKlipper)', 'GCodeWriter.hpp,cpp', 'Klipper 호환 명령어'),
])

doc.add_heading('3.3 제거할 것', level=2)
add_table_with_header(doc, [
    ('제거 대상', '관련 파일', '이유'),
    ('Prusa 계정 로그인', 'UserAccount*, LoginDialog*, WebViewDialog*', 'Prusa 전용'),
    ('PrusaConnect 연동', 'PrusaConnect.cpp, ConnectRequestHandler*', 'Prusa 전용'),
    ('Configuration Sources', 'ConfigWizard.cpp PageUpdateManager', 'Prusa 전용'),
    ('프리셋 자동 업데이트', 'PresetUpdater*, PresetUpdaterWrapper*', 'Prusa 전용'),
    ('기존 프린터 프로파일 36개', 'resources/profiles/', '불필요'),
    ('FDM 냉각 (CoolingBuffer)', 'GCode/CoolingBuffer.*', '레진에 불필요'),
    ('FDM 리트랙션', 'GCode 관련', '레진 도포에 불필요'),
    ('FDM 온도 제어', 'GCode 관련', '레진에 불필요'),
    ('FDM 벽(Perimeter) 정밀 생성', 'PerimeterGenerator, Arachne', '도포는 정밀 벽 불필요'),
])

doc.add_heading('3.4 참고 오픈소스 프로젝트', level=2)
doc.add_paragraph('Open Resin Alliance 프로젝트가 매우 유사한 구조이다:')
add_table_with_header(doc, [
    ('프로젝트', '역할', '당신의 대응', 'GitHub'),
    ('LumenSlicer', 'PrusaSlicer 포크 레진 슬라이서', '커스텀 슬라이서', 'Open-Resin-Alliance/LumenSlicer'),
    ('Odyssey', 'SL1 파일 처리 + Klipper 통신', 'vgui', 'Open-Resin-Alliance/Odyssey'),
    ('Orion', 'Flutter 기반 프린터 GUI', 'vgui', 'Open-Resin-Alliance/Orion'),
])
doc.add_paragraph()
doc.add_paragraph(
    '차이점: Open Resin Alliance는 순수 mSLA(도포 경로 없음)이지만, '
    '출력 포맷(ZIP+PNG), 레이어 이미지 생성, Klipper 통신 구조는 참고할 수 있다.')

doc.add_page_break()

# ============================================================
# 4. 개발 Phase
# ============================================================
doc.add_heading('4. 개발 Phase', level=1)

# Phase 0
doc.add_heading('Phase 0: 빌드 환경 구축', level=2)
add_bold_then_normal(doc, '목표: ', 'PrusaSlicer를 그대로 빌드하고 실행할 수 있는 환경을 만든다.')
doc.add_paragraph('빌드 요구사항: C++17 (MSVC 2019+), CMake 3.13+, Git')
doc.add_paragraph('deps/ 디렉토리의 자동 빌드 스크립트로 28개 의존성 빌드')
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('빌드가 안 되면 아무것도 못 한다. 반드시 이 단계를 먼저 완료할 것.')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

# Phase 1
doc.add_heading('Phase 1: 불필요 기능 제거 + 브랜딩', level=2)
add_bold_then_normal(doc, '목표: ', 'Prusa 전용 기능을 비활성화하고 커스텀 슬라이서 정체성 확립')
doc.add_paragraph('섹션 3.3의 제거 목록 참고. #ifdef나 빌드 옵션으로 비활성화 (코드 삭제보다 안전)')

# Phase 2
doc.add_heading('Phase 2: 멀티 레진 Tool 매핑', level=2)
add_bold_then_normal(doc, '목표: ', 'STL 오브젝트별로 레진(Tool)을 지정할 수 있게 한다.')
doc.add_paragraph()
doc.add_paragraph('PrusaSlicer의 멀티 익스트루더 기능을 활용:')
add_code_block(doc, [
    '기존 멀티 익스트루더:',
    '  오브젝트별로 Extruder 번호 지정 (Extruder 1, 2, ...)',
    '  -> 다른 필라멘트 사용',
    '',
    '하이브리드 슬라이서:',
    '  오브젝트별로 Tool 번호 지정 (T0, T1, ...)',
    '  -> 다른 레진 사용',
    '',
    '  잇몸_모델.stl -> T0 -> 모델레진',
    '  치아_21.stl   -> T1 -> 템프레진',
    '  치아_22.stl   -> T1 -> 템프레진',
])
doc.add_paragraph()
doc.add_paragraph('구현: 기존 extruder 선택 UI/로직을 Tool(레진) 선택으로 리네이밍하여 활용')

# Phase 3
doc.add_heading('Phase 3: 마스크 이미지(PNG) 생성', level=2)
add_bold_then_normal(doc, '목표: ', '각 레이어의 슬라이스 단면을 마스크 이미지(PNG)로 래스터화한다.')
doc.add_paragraph()
doc.add_paragraph('구현 방법:')
for b in [
    'SLA의 RasterBase 클래스를 참고하여 ExPolygon -> PNG 변환',
    '모든 STL의 단면을 하나로 합쳐서 1장의 마스크 이미지 생성',
    '흰색 = 경화 영역, 검은색 = 미경화 영역',
    '해상도는 DLP 프로젝터의 픽셀 해상도에 맞춤',
    'Tool 구분 없이 합쳐진 1장 (경화 시간이 동일하므로)',
]:
    doc.add_paragraph(b, style='List Bullet')

# Phase 4
doc.add_heading('Phase 4: 도포 경로 GCode 생성', level=2)
add_bold_then_normal(doc, '목표: ', '각 레이어에서 Tool별 도포 영역에 대한 GCode를 생성한다.')
doc.add_paragraph()
doc.add_paragraph('레이어별 GCode 구조:')
add_code_block(doc, [
    '; Layer N',
    'T0                        ; 모델레진 노즐 선택',
    'G1 X.. Y.. E.. F..        ; T0 도포 경로 (잇몸 영역)',
    'G1 X.. Y.. E.. F..',
    'T1                        ; 템프레진 노즐 선택',
    'G1 X.. Y.. E.. F..        ; T1 도포 경로 (치아 영역)',
    'G1 X.. Y.. E.. F..',
    '; Blade sweep',
    'G1 X.. F..                ; 블레이드 평탄화',
    '; Expose layer N',
    'DISPLAY_IMAGE layers/00N.png  ; 마스크 이미지 표시 (vgui가 처리)',
    'G4 P8000                  ; 노출 대기 (8초)',
    '; Z move',
    'G1 Z.. F..                ; 다음 레이어 높이',
])
doc.add_paragraph()
doc.add_paragraph(
    '도포 경로는 FDM의 인필 패턴(Rectilinear 등)을 단순화하여 활용 가능. '
    '정밀한 벽(Perimeter) 생성은 불필요하며, 해당 영역을 대략적으로 채우는 수준이면 충분하다.')

# Phase 5
doc.add_heading('Phase 5: 출력 포맷 패키징 (ZIP)', level=2)
add_bold_then_normal(doc, '목표: ', 'GCode + 마스크 PNG + 메타데이터를 ZIP으로 패키징한다.')
doc.add_paragraph()
doc.add_paragraph('SL1Archive를 참고하여 HybridArchiveWriter 작성')
doc.add_paragraph('Zipper 클래스(PrusaSlicer 내장)로 ZIP 생성')
doc.add_paragraph()
doc.add_paragraph('manifest.json 예시:')
add_code_block(doc, [
    '{',
    '  "layer_height": 0.05,',
    '  "exposure_time": 8.0,',
    '  "total_layers": 1200,',
    '  "resolution_x": 3840,',
    '  "resolution_y": 2160,',
    '  "tools": [',
    '    {"id": 0, "name": "Model Resin", "type": "model"},',
    '    {"id": 1, "name": "Temp Resin", "type": "temp"}',
    '  ]',
    '}',
])

# Phase 6
doc.add_heading('Phase 6: vgui 연동 + 실 프린트 테스트', level=2)
add_bold_then_normal(doc, '목표: ', 'vgui가 새 출력 포맷을 읽고 프린팅을 실행할 수 있게 한다.')
doc.add_paragraph('자세한 사항은 섹션 6 참고.')

doc.add_page_break()

# ============================================================
# 5. 주요 수정 파일 목록
# ============================================================
doc.add_heading('5. 주요 수정 파일 목록', level=1)

add_table_with_header(doc, [
    ('파일', '수정 내용', '우선순위'),
    ('src/libslic3r/SLA/RasterBase.*', 'ExPolygon -> PNG 래스터화 (마스크 생성)', '★★★'),
    ('src/libslic3r/Format/SL1.*', '참고하여 HybridArchiveWriter 작성', '★★★'),
    ('src/libslic3r/PrintConfig.hpp', '하이브리드 전용 파라미터 추가', '★★★'),
    ('src/libslic3r/Config.hpp', 'PrinterTechnology에 ptHybrid 추가', '★★★'),
    ('src/libslic3r/GCode.*', '도포 경로 GCode 생성 (Klipper)', '★★☆'),
    ('src/libslic3r/GCode/GCodeWriter.*', '블레이드/마스크 명령어 삽입', '★★☆'),
    ('src/libslic3r/Print.*, PrintObject.*', '하이브리드 파이프라인 분기', '★★☆'),
    ('src/libslic3r/Fill/', '인필 패턴 -> 도포 경로 활용', '★★☆'),
    ('src/libslic3r/Preset.*', '프리셋에 하이브리드 타입 추가', '★★☆'),
    ('resources/profiles/', '커스텀 프린터 프로파일 추가', '★★☆'),
    ('src/slic3r/GUI/MainFrame.*', 'UI 브랜딩', '★☆☆'),
    ('src/slic3r/GUI/ConfigWizard.*', '하이브리드 프린터 선택', '★☆☆'),
    ('src/slic3r/GUI/UserAccount*', '로그인 비활성화', '★☆☆'),
    ('src/slic3r/Utils/PrusaConnect.*', 'PrusaConnect 비활성화', '★☆☆'),
])

doc.add_page_break()

# ============================================================
# 6. vgui (프린터 측) 수정 사항
# ============================================================
doc.add_heading('6. vgui (프린터 측) 수정 사항', level=1)

doc.add_paragraph(
    '슬라이서 출력 포맷이 바뀌면 vgui도 수정이 필요하다. '
    '단, vgui의 핵심 로직(Moonraker 통신, 프로젝터 표시)은 바뀌지 않으며, '
    '파일 파싱 부분만 수정하면 된다.')

doc.add_paragraph()
add_table_with_header(doc, [
    ('바뀌는 것', '안 바뀌는 것'),
    ('ZIP 파일 열기/파싱 방식', 'Moonraker API 호출 로직'),
    ('GCode 파일명/경로 규칙', '모터 제어 흐름'),
    ('PNG 파일명 규칙 (layers/*.png)', '프로젝터 이미지 표시 로직'),
    ('메타데이터 포맷 (manifest.json)', 'UI 레이아웃/디자인'),
    ('미리보기 이미지 위치', 'USB 감지/마운트'),
])

doc.add_paragraph()
doc.add_paragraph('개발 순서 (중요):')
p = doc.add_paragraph()
run = p.add_run('출력 파일 포맷(ZIP 구조, manifest 스펙)을 먼저 확정한 후, 슬라이서와 vgui를 동시에 개발한다.')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('현실적 팁:')
doc.add_paragraph(
    '처음에는 치투슬라이서의 포맷을 그대로 호환시키면 기존 vgui로 바로 테스트할 수 있다. '
    '이후 자체 포맷으로 전환하는 2단계 접근이 개발 속도를 높인다.', style='List Bullet')

doc.add_page_break()

# ============================================================
# 7. 핵심 원칙 + 일정
# ============================================================
doc.add_heading('7. 핵심 원칙 + 일정', level=1)

doc.add_heading('핵심 원칙', level=2)
principles = [
    ('레진 슬라이서가 몸통이다',
     '마스크 이미지(PNG)가 최종 형상을 결정한다. 도포 GCode는 보조 기능이다.'),
    ('기존 기능 위에 추가한다',
     'PrusaSlicer를 포크하고 점진적으로 수정한다. 새로 만들지 않는다.'),
    ('STL을 수정하지 않는다',
     '외부 CAD에서 좌표가 확정된 상태로 들어온다. Boolean 연산 등은 피한다.'),
    ('출력 포맷을 먼저 확정한다',
     '슬라이서(PC)와 vgui(프린터)가 같은 포맷을 쓰지 않으면 의미 없다.'),
    ('한 번에 하나씩',
     '각 Phase를 완료하고 테스트한 후 다음으로 넘어간다.'),
    ('이상적 시나리오로 먼저 개발한다',
     '물리적 문제(레진 혼합, 쏠림 등)는 하드웨어/재료가 해결할 영역이다.'),
]
for i, (t, d) in enumerate(principles, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {t}')
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph(d)

doc.add_heading('권장 일정', level=2)
add_table_with_header(doc, [
    ('기간', 'Phase', '주요 작업'),
    ('1주차', 'Phase 0', '빌드 환경 구축, 정상 빌드/실행 확인'),
    ('2주차', 'Phase 1', '불필요 기능 비활성화, 브랜딩'),
    ('3-4주차', 'Phase 2', '멀티 레진 Tool 매핑 (멀티 익스트루더 활용)'),
    ('5-6주차', 'Phase 3', '마스크 이미지 PNG 생성 (RasterBase 참고)'),
    ('7-8주차', 'Phase 4', '도포 경로 GCode 생성 (인필 패턴 활용)'),
    ('9-10주차', 'Phase 5', 'ZIP 패키징 (SL1Archive 참고)'),
    ('11-12주차', 'Phase 6', 'vgui 연동 + 실 프린트 테스트'),
])

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('-- End --')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# 저장
output_path = r'c:\Users\JoWooHyun\Documents\PrusaSlicer\커스텀_슬라이서_개발계획서_v2.docx'
doc.save(output_path)
print(f'Done: {output_path}')
