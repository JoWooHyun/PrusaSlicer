from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# 기본 스타일 설정
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ============================================================
# 표지
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PrusaSlicer 기반\n커스텀 슬라이서 개발 계획서')
run.font.size = Pt(28)
run.bold = True

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('하이브리드 프린터 (FDM 이동 + LED 경화) 전용 슬라이서')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('기반 소스: PrusaSlicer v2.9.4 (AGPLv3)\n작성일: 2026-03-10')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ============================================================
# 목차
# ============================================================
doc.add_heading('목차', level=1)
toc_items = [
    '1. 프로젝트 개요',
    '2. PrusaSlicer 아키텍처 분석',
    '   2.1 오픈소스 라이브러리 구성',
    '   2.2 슬라이싱 파이프라인',
    '   2.3 핵심 클래스 구조',
    '   2.4 디렉토리 구조',
    '3. 커스텀 슬라이서 개발 계획',
    '   Phase 0: 빌드 환경 구축',
    '   Phase 1: 브랜딩 + 불필요 기능 제거',
    '   Phase 2: PrinterTechnology 추가',
    '   Phase 3: 출력 포맷 구현',
    '   Phase 4: 슬라이싱 파이프라인 수정',
    '   Phase 5: 프린터 프로파일 생성',
    '4. 주요 수정 파일 목록',
    '5. 핵심 원칙',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

doc.add_page_break()

# ============================================================
# 1. 프로젝트 개요
# ============================================================
doc.add_heading('1. 프로젝트 개요', level=1)

doc.add_heading('프린터 특수성', level=2)
doc.add_paragraph(
    '이 프로젝트는 FDM과 SLA를 결합한 독자적인 하이브리드 방식의 3D 프린터를 위한 '
    '커스텀 슬라이서를 개발하는 것이다. PrusaSlicer(v2.9.4) 소스코드를 기반으로 한다.'
)

bullets = [
    'FDM처럼 XYZ축 이동을 하지만, 필라멘트 대신 고점도 레진을 도포',
    '도포된 레진에 LED를 조사하여 경화하는 하이브리드 방식',
    '따라서 슬라이서는 GCode(도포 경로) + 레이어 이미지(LED 패턴) 둘 다 생성해야 함',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('사용 워크플로우', level=2)
steps = [
    'PC(Windows)에서 커스텀 슬라이서로 3D 모델을 슬라이싱',
    '슬라이싱 결과 파일을 USB에 저장',
    'USB를 프린터(Raspberry Pi)에 삽입',
    'vgui(PySide6 GUI)가 파일을 읽어서 프린팅 실행',
    '   - GCode → Moonraker API → Klipper로 모터 제어 (레진 도포)',
    '   - 레이어 이미지 → 프로젝터로 LED 조사 (레진 경화)',
]
for i, s in enumerate(steps, 1):
    if s.startswith('   '):
        doc.add_paragraph(s.strip(), style='List Bullet 2')
    else:
        doc.add_paragraph(f'{i}. {s}')

p = doc.add_paragraph()
run = p.add_run('※ 슬라이서는 PC 전용 데스크탑 앱이며, 프린터와 직접 네트워크 통신하지 않음 (오프라인 USB 전달)')
run.bold = True

doc.add_heading('하드웨어 환경', level=2)
hw_table = doc.add_table(rows=6, cols=2, style='Light Shading Accent 1')
hw_data = [
    ('구분', '사양'),
    ('보드', 'BigTreeTech Manta M8P 2.0'),
    ('컴퓨트', 'Raspberry Pi CM4'),
    ('펌웨어', 'Klipper + Moonraker'),
    ('프린터 GUI', 'vgui (PySide6, 자체 개발)'),
    ('현재 슬라이서', '치투슬라이서 (교체 예정)'),
]
for i, (k, v) in enumerate(hw_data):
    hw_table.rows[i].cells[0].text = k
    hw_table.rows[i].cells[1].text = v
    if i == 0:
        for cell in hw_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_heading('개발 목표', level=2)
goals = [
    '슬라이싱 엔진을 레진 도포 방식에 맞게 수정',
    'Klipper 호환 GCode 생성 (레진 도포 경로용)',
    '레이어별 LED 조사 이미지(PNG) 생성 기능 추가',
    '출력 파일 포맷 설계 (GCode + 레이어 이미지 패키지, vgui 호환)',
    'UI/UX 커스터마이징',
    '독자적 기능 추가',
]
for i, g in enumerate(goals, 1):
    doc.add_paragraph(f'{i}. {g}')

doc.add_page_break()

# ============================================================
# 2. PrusaSlicer 아키텍처 분석
# ============================================================
doc.add_heading('2. PrusaSlicer 아키텍처 분석', level=1)

doc.add_heading('2.1 오픈소스 라이브러리 구성', level=2)
doc.add_paragraph(
    'PrusaSlicer는 28개의 외부 오픈소스 라이브러리를 조합하여 만들어졌다. '
    'PrusaSlicer 자체가 만든 것은 이 라이브러리들을 하나의 슬라이싱 파이프라인으로 '
    '통합한 "접착제" + "비즈니스 로직" 부분이다.'
)

lib_table = doc.add_table(rows=11, cols=3, style='Light Shading Accent 1')
lib_data = [
    ('라이브러리', '용도', 'GitHub'),
    ('Clipper / Clipper2', '2D 폴리곤 Boolean 연산 (union, diff, intersection)', 'AngusJohnson/Clipper2'),
    ('CGAL', '3D 메시 Boolean, self-intersection 감지/복구', 'CGAL/cgal'),
    ('libigl', '메시 self-union, winding number', 'libigl/libigl'),
    ('ADMesh', 'STL 로드/복구 (구멍 메우기, 법선 수정)', 'admesh/admesh'),
    ('Eigen3', '선형대수, 행렬/벡터 연산', 'eigen/eigen'),
    ('TBB', '멀티스레드 병렬 처리', 'oneapi-src/oneTBB'),
    ('Boost', '파일시스템, 로깅, 스레드, 직렬화', 'boostorg/boost'),
    ('OpenVDB', '3D 복셀 처리', 'AcademySoftwareFoundation/openvdb'),
    ('wxWidgets', 'GUI 프레임워크', 'wxWidgets/wxWidgets'),
    ('NLopt', '수치 최적화', 'stevengj/nlopt'),
]
for i, (name, purpose, github) in enumerate(lib_data):
    lib_table.rows[i].cells[0].text = name
    lib_table.rows[i].cells[1].text = purpose
    lib_table.rows[i].cells[2].text = github
    if i == 0:
        for cell in lib_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('PrusaSlicer의 핵심 가치: ')
run.bold = True
p.add_run('알고리즘 자체가 아니라, 28개 라이브러리를 하나의 매끄러운 파이프라인으로 통합한 것. '
          '"어떤 순서로, 어떤 조건에서 적용할지"가 PrusaSlicer의 노하우.')

doc.add_heading('2.2 슬라이싱 파이프라인', level=2)

pipeline_steps = [
    ('STL/OBJ/3MF 로드', 'Model → TriangleMesh 변환'),
    ('posSlice', 'TriangleMeshSlicer::slice_mesh_ex() → 3D 메시를 Z 평면으로 절단 → ExPolygons 생성'),
    ('posPerimeters', 'PerimeterGenerator::process() → Arachne 알고리즘으로 벽(Perimeter) 생성'),
    ('posPrepareInfill', '서피스 감지, 브릿지 감지, 쉘 결합'),
    ('posInfill', 'Fill::fill_surface() → Rectilinear/Honeycomb/Gyroid 등 인필 패턴 적용'),
    ('posIroning', '상단 면 아이로닝 (선택적)'),
    ('posSupportMaterial', 'TreeSupport/SupportMaterial → 서포트 구조 생성'),
    ('posEstimateCurledExtrusions', '말림(Curling) 추정'),
    ('후처리', 'WipeTower / Skirt / Brim 생성'),
    ('GCode 출력', 'GCodeGenerator::do_export() → CoolingBuffer, SpiralVase, PressureEqualizer 등 후처리 → 최종 GCode'),
]

pipe_table = doc.add_table(rows=len(pipeline_steps)+1, cols=2, style='Light Shading Accent 1')
pipe_table.rows[0].cells[0].text = '단계'
pipe_table.rows[0].cells[1].text = '설명'
for cell in pipe_table.rows[0].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
for i, (stage, desc) in enumerate(pipeline_steps, 1):
    pipe_table.rows[i].cells[0].text = stage
    pipe_table.rows[i].cells[1].text = desc

doc.add_heading('2.3 핵심 클래스 구조', level=2)

class_items = [
    'Print — 프린트 작업 관리 (최상위)',
    '  ├─ PrintObject[] — 프린트 객체',
    '  │   ├─ Layer[] — 레이어',
    '  │   │   └─ LayerRegion[] — 영역 (perimeters, fills, slices)',
    '  │   ├─ SupportLayer[] — 서포트 레이어',
    '  │   └─ PrintObjectRegions → PrintRegion[]',
    '  ├─ GCodeGenerator — GCode 생성',
    '  │   ├─ GCodeWriter — 명령어 작성',
    '  │   ├─ CoolingBuffer — 냉각 제어',
    '  │   ├─ Seams::Placer — 이음새 위치',
    '  │   └─ AvoidCrossingPerimeters — 경로 최적화',
    '  └─ PrintConfig — 300+ 설정 파라미터',
]
for item in class_items:
    p = doc.add_paragraph(item)
    p.style = doc.styles['No Spacing']
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

doc.add_heading('2.4 디렉토리 구조', level=2)

dir_items = [
    ('src/libslic3r/', '핵심 슬라이싱 엔진 (C++ 라이브러리)'),
    ('src/libslic3r/GCode/', 'GCode 생성/처리'),
    ('src/libslic3r/Fill/', '인필 패턴 (8+ 종류)'),
    ('src/libslic3r/Support/', '서포트 생성'),
    ('src/libslic3r/SLA/', 'SLA(레진) 프린터 처리 ★ 레이어 이미지 참고'),
    ('src/libslic3r/Format/', '파일 형식 (STL, OBJ, 3MF, SL1 등) ★ 출력 포맷 참고'),
    ('src/libslic3r/Arachne/', '벽(Perimeter) 생성 알고리즘'),
    ('src/slic3r/GUI/', 'wxWidgets 기반 GUI'),
    ('src/clipper/', '다각형 연산 (Clipper 번들)'),
    ('deps/', '외부 의존성 빌드 (28개 라이브러리)'),
    ('resources/profiles/', '프린터 프리셋'),
    ('tests/', '단위 테스트 (Catch2)'),
]

dir_table = doc.add_table(rows=len(dir_items)+1, cols=2, style='Light Shading Accent 1')
dir_table.rows[0].cells[0].text = '경로'
dir_table.rows[0].cells[1].text = '설명'
for cell in dir_table.rows[0].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
for i, (path, desc) in enumerate(dir_items, 1):
    dir_table.rows[i].cells[0].text = path
    dir_table.rows[i].cells[1].text = desc

doc.add_page_break()

# ============================================================
# 3. 커스텀 슬라이서 개발 계획
# ============================================================
doc.add_heading('3. 커스텀 슬라이서 개발 계획', level=1)

# Phase 0
doc.add_heading('Phase 0: 빌드 환경 구축', level=2)
p = doc.add_paragraph()
run = p.add_run('목표: ')
run.bold = True
p.add_run('PrusaSlicer를 그대로 빌드하고 실행할 수 있는 환경을 만든다. 이것이 기준점(baseline)이 된다.')

doc.add_paragraph('빌드 요구사항:', style='List Bullet')
build_reqs = [
    'C++17 호환 컴파일러 (MSVC 2019+)',
    'CMake 3.13+',
    'Git',
    '28개 외부 의존성 (deps/ 디렉토리의 자동 빌드 스크립트 사용)',
]
for r in build_reqs:
    doc.add_paragraph(r, style='List Bullet 2')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('빌드가 안 되면 아무것도 못 한다. 반드시 이 단계를 먼저 완료할 것.')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

# Phase 1
doc.add_heading('Phase 1: 브랜딩 + 불필요 기능 제거', level=2)
p = doc.add_paragraph()
run = p.add_run('목표: ')
run.bold = True
p.add_run('Prusa 전용 기능을 비활성화하고, 커스텀 슬라이서로서의 정체성을 확립한다.')

remove_table = doc.add_table(rows=6, cols=2, style='Light Shading Accent 1')
remove_data = [
    ('제거/비활성화 대상', '관련 파일'),
    ('Prusa 계정 로그인', 'UserAccount*, LoginDialog*, WebViewDialog*'),
    ('PrusaConnect 연동', 'PrusaConnect.cpp, ConnectRequestHandler*'),
    ('Configuration Sources', 'ConfigWizard.cpp의 PageUpdateManager'),
    ('프리셋 자동 업데이트', 'PresetUpdater*, PresetUpdaterWrapper*'),
    ('기존 프린터 프로파일', 'resources/profiles/ (36개 벤더)'),
]
for i, (target, files) in enumerate(remove_data):
    remove_table.rows[i].cells[0].text = target
    remove_table.rows[i].cells[1].text = files
    if i == 0:
        for cell in remove_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('주의: ')
run.bold = True
p.add_run('코드를 삭제하는 것이 아니라 #ifdef나 빌드 옵션으로 비활성화하는 것이 안전하다.')

# Phase 2
doc.add_heading('Phase 2: 새로운 PrinterTechnology 추가 (핵심)', level=2)
p = doc.add_paragraph()
run = p.add_run('목표: ')
run.bold = True
p.add_run('하이브리드 프린터 타입을 PrusaSlicer 아키텍처에 등록한다.')

doc.add_paragraph('현재 정의 (src/libslic3r/Config.hpp:243):')
p = doc.add_paragraph()
p.style = doc.styles['No Spacing']
code_lines = [
    'enum PrinterTechnology : unsigned char {',
    '    ptFFF,      // 필라멘트 (FDM)',
    '    ptSLA,      // 레진 (SLA)',
    '    ptUnknown,',
    '    ptAny',
    '};',
]
for line in code_lines:
    p = doc.add_paragraph(line)
    p.style = doc.styles['No Spacing']
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph('수정 후:')
code_lines2 = [
    'enum PrinterTechnology : unsigned char {',
    '    ptFFF,',
    '    ptSLA,',
    '    ptHybrid,   // ★ 새로 추가: FDM 이동 + LED 경화',
    '    ptUnknown,',
    '    ptAny',
    '};',
]
for line in code_lines2:
    p = doc.add_paragraph(line)
    p.style = doc.styles['No Spacing']
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('이 enum 하나가 PrusaSlicer 전체에서 분기 조건으로 사용된다. ')
run.bold = True
p.add_run('이걸 추가하면 슬라이싱 엔진, GCode 생성, UI, 설정 등에서 하이브리드 전용 경로를 만들 수 있다.')

# Phase 3
doc.add_heading('Phase 3: 출력 포맷 구현', level=2)
p = doc.add_paragraph()
run = p.add_run('목표: ')
run.bold = True
p.add_run('GCode + 레이어 PNG를 하나의 ZIP 파일로 패키징하는 출력 포맷을 만든다.')

doc.add_paragraph('SL1Archive (src/libslic3r/Format/SL1.hpp)가 가장 가까운 참고 대상이다. '
                  'SL1은 이미 ZIP 안에 레이어 PNG + 메타데이터를 넣는 구조.')

doc.add_paragraph('출력 파일 구조:')
format_lines = [
    'output.zip',
    '├── run.gcode          ← FDM식 레진 도포 경로 (Klipper용)',
    '├── layer_0001.png     ← LED 경화 패턴',
    '├── layer_0002.png',
    '├── ...',
    '├── preview.png        ← 미리보기 썸네일',
    '└── manifest.json      ← 메타데이터 (레이어 높이, 노출 시간 등)',
]
for line in format_lines:
    p = doc.add_paragraph(line)
    p.style = doc.styles['No Spacing']
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph('구현 방법:')
impl_steps = [
    'SLAArchiveWriter를 참고하여 HybridArchiveWriter 클래스 작성',
    'Zipper 클래스(이미 내장됨)로 ZIP 패키징',
    'SLAArchiveFormatRegistry에 새 포맷 등록',
]
for i, s in enumerate(impl_steps, 1):
    doc.add_paragraph(f'{i}. {s}')

# Phase 4
doc.add_heading('Phase 4: 슬라이싱 파이프라인 수정', level=2)
p = doc.add_paragraph()
run = p.add_run('목표: ')
run.bold = True
p.add_run('FDM 파이프라인을 최대한 재활용하면서, 레이어 이미지 생성을 추가한다.')

doc.add_paragraph('기존 FDM 파이프라인:')
p = doc.add_paragraph('STL → Slice → Perimeters → Infill → Support → GCode')
p.style = doc.styles['No Spacing']
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph('하이브리드 파이프라인:')
hybrid_lines = [
    'STL → Slice → Perimeters → Infill → Support',
    '    → GCode (도포 경로)          ← FDM 부분 재활용',
    '    → Layer Images (LED 패턴)    ★ 새로 추가',
    '    → ZIP 패키징                 ★ 새로 추가',
]
for line in hybrid_lines:
    p = doc.add_paragraph(line)
    p.style = doc.styles['No Spacing']
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('FDM의 슬라이싱 + GCode 생성은 거의 그대로 쓸 수 있다.')
run.bold = True

doc.add_paragraph('추가 구현 필요 항목:')
additions = [
    '각 레이어의 ExPolygon을 PNG로 래스터화 (SLA의 RasterBase 참고)',
    'GCode + PNG를 하나의 ZIP으로 묶기 (Zipper 클래스 활용)',
]
for a in additions:
    doc.add_paragraph(a, style='List Bullet')

# Phase 5
doc.add_heading('Phase 5: 프린터 프로파일 생성', level=2)
p = doc.add_paragraph()
run = p.add_run('목표: ')
run.bold = True
p.add_run('하이브리드 프린터의 설정 파라미터와 프로파일을 정의한다.')

doc.add_paragraph('resources/profiles/에 프린터 프로파일 추가:')
profile_lines = [
    'resources/profiles/MyPrinter/',
    '├── MyPrinter.ini         ← 프린터 기본 설정',
    '├── MyPrinter_hybrid.ini  ← 하이브리드 파라미터',
    '└── ...',
]
for line in profile_lines:
    p = doc.add_paragraph(line)
    p.style = doc.styles['No Spacing']
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph('PrintConfig.hpp에 새 파라미터 추가 예시:')
params = [
    '레진 도포 속도 (resin_dispense_speed)',
    'LED 노출 시간 (led_exposure_time)',
    '레이어 이미지 해상도 (layer_image_resolution)',
    'LED 출력 세기 (led_power)',
    '첫 레이어 노출 시간 (first_layer_exposure_time)',
]
for param in params:
    doc.add_paragraph(param, style='List Bullet')

doc.add_page_break()

# ============================================================
# 4. 주요 수정 파일 목록
# ============================================================
doc.add_heading('4. 주요 수정 파일 목록', level=1)

files_table = doc.add_table(rows=16, cols=3, style='Light Shading Accent 1')
files_data = [
    ('파일', '수정 내용', '우선순위'),
    ('src/libslic3r/Config.hpp', 'PrinterTechnology에 ptHybrid 추가', '★★★'),
    ('src/libslic3r/PrintConfig.hpp', '하이브리드 전용 파라미터 추가 (300+)', '★★★'),
    ('src/libslic3r/Format/SL1.hpp,cpp', '참고하여 HybridArchive 작성', '★★★'),
    ('src/libslic3r/SLA/RasterBase.hpp,cpp', 'ExPolygon → PNG 래스터화 참고', '★★★'),
    ('src/libslic3r/GCode.hpp,cpp', 'GCode 생성 커스터마이징 (Klipper)', '★★☆'),
    ('src/libslic3r/GCode/GCodeWriter.hpp,cpp', 'GCode 명령어 작성 수정', '★★☆'),
    ('src/libslic3r/Print.hpp,cpp', '프린트 프로세스에 하이브리드 분기', '★★☆'),
    ('src/libslic3r/PrintObject.cpp', '객체별 슬라이싱 분기', '★★☆'),
    ('src/libslic3r/Preset.hpp,cpp', '프리셋 관리에 하이브리드 타입 추가', '★★☆'),
    ('src/slic3r/GUI/ConfigWizard.cpp', '위자드에서 하이브리드 프린터 선택', '★☆☆'),
    ('src/slic3r/GUI/MainFrame.cpp', 'UI 브랜딩 수정', '★☆☆'),
    ('src/slic3r/GUI/UserAccount*', '로그인 기능 비활성화', '★☆☆'),
    ('src/slic3r/Utils/PrusaConnect.cpp', 'PrusaConnect 비활성화', '★☆☆'),
    ('resources/profiles/', '커스텀 프린터 프로파일 추가', '★★☆'),
]
for i, row_data in enumerate(files_data):
    for j, cell_text in enumerate(row_data):
        files_table.rows[i].cells[j].text = cell_text
    if i == 0:
        for cell in files_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_page_break()

# ============================================================
# 5. 핵심 원칙
# ============================================================
doc.add_heading('5. 핵심 원칙', level=1)

principles = [
    ('처음부터 새로 만들지 마라',
     'PrusaSlicer를 포크(fork)하고, 점진적으로 수정한다. 한 번에 하나씩 바꾼다.'),
    ('FDM 파이프라인을 최대한 살려라',
     '슬라이싱 → GCode 부분은 거의 그대로 쓸 수 있다. 바퀴를 재발명하지 않는다.'),
    ('SLA 코드에서 레이어 이미지 생성 로직을 가져와라',
     'RasterBase, SL1Archive가 가장 가까운 참고 대상이다.'),
    ('출력 포맷을 vgui와 먼저 맞춰라',
     '슬라이서와 프린터 양쪽이 같은 포맷을 쓰지 않으면 의미 없다. 포맷 설계를 먼저 확정하라.'),
    ('한 번에 하나씩',
     '전체를 동시에 바꾸면 디버깅이 불가능해진다. 각 Phase를 완료하고 테스트한 후 다음으로 넘어간다.'),
    ('빌드-실행-테스트 사이클을 짧게 유지하라',
     '수정할 때마다 빌드하고 실행해서 확인한다. 큰 변경을 한꺼번에 하지 않는다.'),
]

for i, (title, desc) in enumerate(principles, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {title}')
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph(desc)

doc.add_paragraph()

# 실제 작업 일정
doc.add_heading('권장 작업 일정', level=2)

schedule_table = doc.add_table(rows=7, cols=3, style='Light Shading Accent 1')
schedule_data = [
    ('기간', 'Phase', '주요 작업'),
    ('1주차', 'Phase 0', '빌드 환경 구축, 정상 빌드/실행 확인'),
    ('2주차', 'Phase 1', '코드 구조 파악, 불필요 기능 비활성화, 브랜딩'),
    ('3-4주차', 'Phase 2', 'PrinterTechnology 추가 + 프린터 프로파일, gcfKlipper 고정'),
    ('5-6주차', 'Phase 3', '출력 포맷 구현 (ZIP = GCode + Layer PNGs), SL1Archive 참고'),
    ('7-8주차', 'Phase 4', '레이어 이미지 생성 (ExPolygon → PNG), RasterBase 참고'),
    ('이후', 'Phase 5+', '슬라이싱 파라미터 튜닝, vgui 연동 테스트, UI 커스터마이징'),
]
for i, (period, phase, work) in enumerate(schedule_data):
    schedule_table.rows[i].cells[0].text = period
    schedule_table.rows[i].cells[1].text = phase
    schedule_table.rows[i].cells[2].text = work
    if i == 0:
        for cell in schedule_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()
doc.add_paragraph()

# 마무리
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 끝 —')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# 저장
output_path = r'c:\Users\JoWooHyun\Documents\PrusaSlicer\커스텀_슬라이서_개발계획서.docx'
doc.save(output_path)
print(f'문서 저장 완료: {output_path}')
