from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def add_code(doc, lines):
    for line in lines:
        p = doc.add_paragraph(line)
        p.style = doc.styles['No Spacing']
        for r in p.runs:
            r.font.name = 'Consolas'
            r.font.size = Pt(9)

def add_table(doc, data):
    t = doc.add_table(rows=len(data), cols=len(data[0]), style='Light Shading Accent 1')
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            t.rows[i].cells[j].text = str(cell)
        if i == 0:
            for c in t.rows[i].cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.bold = True
    return t

def bold_then(doc, b, n):
    p = doc.add_paragraph()
    r = p.add_run(b)
    r.bold = True
    p.add_run(n)
    return p

# ============================================================
# 표지
# ============================================================
for _ in range(6):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('하이브리드 멀티레진 3D 프린터\n커스텀 슬라이서 개발 계획서 v3')
r.font.size = Pt(26)
r.bold = True
doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Top-Down | FDM 도포 + 블레이드 평탄화 + DLP 마스크 경화 | 멀티레진')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()
doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run('기반: PrusaSlicer v2.9.4 (AGPLv3)\n참고: hybrid_resin_slicer_guideline.docx\n작성일: 2026-03-10 | v3')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
doc.add_page_break()

# ============================================================
# 목차
# ============================================================
doc.add_heading('목차', level=1)
toc = [
    '1. 프로젝트 개요',
    '   1.1 프린터 방식',
    '   1.2 레이어 공정 시퀀스 (5단계)',
    '   1.3 멀티 레진 + 재료 할당 구조',
    '   1.4 사용 워크플로우',
    '   1.5 하드웨어 환경',
    '2. 슬라이서 설계 방향',
    '   2.1 핵심 관점: 레진 슬라이서가 몸통',
    '   2.2 슬라이싱 파이프라인',
    '   2.3 출력 데이터 + manifest.json 스펙',
    '   2.4 입력 모델 형식',
    '   2.5 설계 원칙',
    '3. 하드웨어 제어 파라미터',
    '   3.1 듀얼 노즐 오프셋',
    '   3.2 디스펜서 제어 (공기압)',
    '   3.3 블레이드 제어',
    '   3.4 LED/프로젝터 제어 + Flat Field 보정',
    '4. PrusaSlicer 활용 전략',
    '   4.1 SLA에서 가져올 것 (몸통)',
    '   4.2 FDM에서 가져올 것 (추가)',
    '   4.3 제거할 것 (FDM/Prusa 전용)',
    '   4.4 참고 오픈소스',
    '5. 개발 Phase',
    '6. 주요 수정 파일 + 코드 탐색 가이드',
    '7. vgui 수정 사항',
    '8. 핵심 원칙 + 일정',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
doc.add_page_break()

# ============================================================
# 1. 프로젝트 개요
# ============================================================
doc.add_heading('1. 프로젝트 개요', level=1)

doc.add_heading('1.1 프린터 방식', level=2)
doc.add_paragraph(
    'Top-Down 방식의 하이브리드 멀티레진 3D 프린터를 위한 커스텀 슬라이서를 개발한다. '
    'PrusaSlicer(v2.9.4) 소스코드를 기반으로 하되, 레진 도포 시스템에 맞게 수정한다.')
doc.add_paragraph('프린터는 세 가지 기술을 결합한다:')
for b in [
    'FDM 방식의 도포 경로 (XYZ 이동 + 공기압 디스펜서로 레진 공급)',
    '블레이드 평탄화 (도포된 레진을 균일하게 펼침)',
    'DLP 마스크 경화 (LED 프로젝터로 선택적 경화)',
]:
    doc.add_paragraph(b, style='List Bullet')
doc.add_paragraph()
bold_then(doc, '핵심: ', '도포 경로(GCode)는 재료 공급만 담당. 최종 형상의 정밀도는 마스크 이미지(PNG)가 결정.')

doc.add_heading('1.2 레이어 공정 시퀀스 (5단계)', level=2)
doc.add_paragraph('각 레이어는 반드시 다음 순서로 처리된다. 이 시퀀스는 슬라이서, vgui, 펌웨어 모두 동일하게 따라야 한다.')
doc.add_paragraph()
add_code(doc, [
    'for each layer N:',
    '    1. T0 deposition     # Tool0 노즐로 모델레진 도포 (GCode)',
    '    2. T1 deposition     # Tool1 노즐로 템프레진 도포 (GCode)',
    '    3. Blade leveling    # 블레이드 스윕으로 평탄화 (GCode)',
    '    4. LED exposure      # 마스크 이미지 조사 + 경화 대기 (PNG + G4)',
    '    5. Z lift            # 다음 레이어 높이로 이동 (GCode)',
])
doc.add_paragraph()

add_table(doc, [
    ('단계', '동작', '출력 데이터', '비고'),
    ('1. T0 도포', '모델레진 영역에 T0 노즐로 도포', 'GCode (G1 X Y E)', '노즐 오프셋 보정 적용'),
    ('2. T1 도포', '템프레진 영역에 T1 노즐로 도포', 'GCode (G1 X Y E)', '노즐 오프셋 보정 적용'),
    ('3. 블레이드', '레진 표면을 고르게 평탄화', 'GCode (블레이드)', '속도/방향/횟수 설정'),
    ('4. 마스크 노광', '프로젝터가 마스크 이미지 조사', 'PNG (1장, 합침)', '노출 시간 대기'),
    ('5. Z 이동', '다음 레이어 높이로 이동', 'GCode (G1 Z)', '레이어 높이만큼'),
])

doc.add_paragraph()
bold_then(doc, '가정 (이상적 시나리오): ',
    '두 레진의 경화 시간 동일, 블레이드 평탄화 시 레진 혼합 없음, 도포 영역 정확히 분리됨. '
    '물리적 문제는 하드웨어/재료가 해결할 영역이며, 소프트웨어는 이상적으로 동작한다고 전제.')

doc.add_heading('1.3 멀티 레진 + 재료 할당 구조', level=2)
doc.add_paragraph('각 STL 객체는 다음 체계로 재료에 매핑된다:')
doc.add_paragraph()
add_code(doc, [
    'Object (STL 파일)',
    '  -> Material (재료 ID)',
    '    -> Tool (도포 노즐)',
    '      -> Nozzle (물리적 하드웨어)',
    '',
    '예시 (치과 모델):',
    '  잇몸_모델.stl  -> Material 0 (모델레진) -> Tool 0 -> Nozzle 0',
    '  치아_21.stl    -> Material 1 (템프레진) -> Tool 1 -> Nozzle 1',
    '  치아_22.stl    -> Material 1 (템프레진) -> Tool 1 -> Nozzle 1',
])
doc.add_paragraph()
doc.add_paragraph(
    '모델들은 외부 CAD/치과 소프트웨어(exocad 등)에서 좌표가 이미 정해진 상태로 준비된다. '
    '슬라이서는 STL을 수정하지 않고 그대로 슬라이스한다. '
    'STL 파일들은 서로 겹칠 수 있다 (예: 잇몸 구멍 안에 치아가 들어감).')
doc.add_paragraph()
doc.add_paragraph('레진 혼합 정책:')
for b in [
    '레진 혼합은 허용된다',
    '경화되지 않은 혼합 레진은 재사용하지 않고 폐기',
    '목표는 완벽한 재료 분리가 아닌 멀티 레진 출력 기능 구현',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('1.4 사용 워크플로우', level=2)
add_code(doc, [
    '[외부 CAD]           [PC]                      [USB]          [프린터]',
    'exocad 등에서     커스텀 슬라이서에서        USB에 저장      vgui가 파일을',
    'STL 파일 준비 -> 로드 + Material 지정  -> ZIP 파일  ->  읽어서 프린팅',
    '(좌표 확정)       슬라이스                               실행',
])
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('슬라이서는 PC 전용 데스크탑 앱. 프린터와 직접 통신하지 않음 (오프라인 USB 전달)')
r.bold = True

doc.add_heading('1.5 하드웨어 환경', level=2)
add_table(doc, [
    ('구분', '사양'),
    ('보드', 'BigTreeTech Manta M8P 2.0'),
    ('컴퓨트', 'Raspberry Pi CM4'),
    ('펌웨어', 'Klipper + Moonraker'),
    ('프린터 GUI', 'vgui (PySide6, 자체 개발)'),
    ('프로젝터', 'DLP LED 프로젝터 (마스크 경화용)'),
    ('도포', '공기압 디스펜서 + 멀티 노즐 (T0, T1)'),
    ('블레이드', '레진 평탄화용 (편도/왕복 설정 가능)'),
    ('현재 슬라이서', '치투슬라이서 (교체 예정)'),
])

doc.add_page_break()

# ============================================================
# 2. 슬라이서 설계 방향
# ============================================================
doc.add_heading('2. 슬라이서 설계 방향', level=1)

doc.add_heading('2.1 핵심 관점: 레진 슬라이서가 몸통', level=2)
p = doc.add_paragraph()
r = p.add_run('중요: ')
r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run('이 슬라이서의 몸통은 FDM이 아니라 레진(SLA) 슬라이서이다.')
doc.add_paragraph()
add_code(doc, [
    '잘못된 접근: FDM 슬라이서가 몸통 + SLA 기능 추가',
    '올바른 접근: 레진 슬라이서가 몸통 + FDM 도포 기능 추가',
    '',
    '┌─────────────────────────────────────────────┐',
    '│  레진 슬라이서 (몸통)                         │',
    '│  ├─ 3D 뷰어 + STL/3MF 배치                   │',
    '│  ├─ 레이어 슬라이싱 (Z 절단 -> ExPolygon)     │',
    '│  ├─ 마스크 이미지 생성 (PNG)    <- 핵심       │',
    '│  ├─ 서포트 생성 (레진 방식)                    │',
    '│  ├─ 미리보기                                  │',
    '│  │                                            │',
    '│  ├─ [추가] 재료 할당 (Object->Material->Tool) │',
    '│  ├─ [추가] 도포 경로 GCode (공기압 디스펜서)  │',
    '│  ├─ [추가] 듀얼 노즐 오프셋 보정              │',
    '│  ├─ [추가] 블레이드 평탄화 GCode              │',
    '│  ├─ [추가] Flat Field 보정                    │',
    '│  └─ [추가] ZIP 패키징                         │',
    '└─────────────────────────────────────────────┘',
])

doc.add_heading('2.2 슬라이싱 파이프라인', level=2)
add_code(doc, [
    'STL/3MF 여러 개 로드 (외부 CAD에서 좌표 확정 상태)',
    '    |',
    '    v',
    'Object별 Material/Tool 지정 (Object->Material->Tool->Nozzle)',
    '    |',
    '    v',
    'Z 슬라이싱 (레이어 높이별로 절단 -> ExPolygons)',
    '    |',
    '    +---> 마스크 이미지 생성 (전체 합쳐서 1장 PNG)',
    '    |     (+ Flat Field 보정 적용)',
    '    |',
    '    +---> Tool별 도포 영역 -> 도포 경로 GCode 생성',
    '    |     (+ 듀얼 노즐 오프셋 보정)',
    '    |     (+ 디스펜서 시작/정지 지연 보정)',
    '    |',
    '    +---> 블레이드/Z이동 GCode 삽입',
    '    |',
    '    v',
    'ZIP 패키징 (run.gcode + layers/*.png + manifest.json)',
])

doc.add_heading('2.3 출력 데이터 + manifest.json 스펙', level=2)
doc.add_paragraph('슬라이서 출력:')
doc.add_paragraph('1. GCode: 도포 경로 + 블레이드 + Z이동 (Klipper 호환)')
doc.add_paragraph('2. PNG 마스크 이미지: 레이어별 경화 패턴 (재료 구분 없이 합쳐진 1장)')
doc.add_paragraph()
doc.add_paragraph('파일 구조:')
add_code(doc, [
    'print_job.zip',
    '├── run.gcode                <- 전체 도포/블레이드/Z이동 명령',
    '├── layers/',
    '│   ├── 0000.png             <- 레이어 0 마스크',
    '│   ├── 0001.png',
    '│   └── ...',
    '├── preview.png              <- 미리보기 썸네일',
    '├── flat_field_mask.png      <- Flat Field 보정 마스크 (선택)',
    '└── manifest.json            <- 메타데이터',
])

doc.add_paragraph()
doc.add_paragraph('manifest.json 상세 스펙:')
add_code(doc, [
    '{',
    '  "version": "1.0",',
    '  "slicer": "HybridResinSlicer",',
    '',
    '  "layer_height": 0.05,',
    '  "total_layers": 1200,',
    '  "first_layer_height": 0.1,',
    '',
    '  "exposure_time": 2.5,',
    '  "first_layer_exposure_time": 8.0,',
    '',
    '  "resolution_x": 2560,',
    '  "resolution_y": 1440,',
    '  "pixel_size_um": 50,',
    '',
    '  "tools": [',
    '    {',
    '      "id": 0,',
    '      "name": "Model Resin",',
    '      "type": "model",',
    '      "nozzle_offset_x": 0.0,',
    '      "nozzle_offset_y": 0.0,',
    '      "nozzle_diameter": 0.4,',
    '      "dispense_pressure": 300,',
    '      "dispense_start_delay_ms": 50,',
    '      "dispense_stop_delay_ms": 30',
    '    },',
    '    {',
    '      "id": 1,',
    '      "name": "Temp Resin",',
    '      "type": "temp",',
    '      "nozzle_offset_x": 35.0,',
    '      "nozzle_offset_y": 0.0,',
    '      "nozzle_diameter": 0.4,',
    '      "dispense_pressure": 280,',
    '      "dispense_start_delay_ms": 60,',
    '      "dispense_stop_delay_ms": 40',
    '    }',
    '  ],',
    '',
    '  "blade": {',
    '    "speed": 10.0,',
    '    "direction": "one_way",',
    '    "pass_count": 1,',
    '    "height": 0.0',
    '  },',
    '',
    '  "flat_field_enabled": false,',
    '  "flat_field_file": "flat_field_mask.png"',
    '}',
])

doc.add_heading('2.4 입력 모델 형식', level=2)
doc.add_paragraph(
    '입력 모델은 STL 또는 3MF를 지원한다. '
    '여러 STL을 동시에 불러와 멀티 재료 출력이 가능하다. '
    '3MF의 경우 하나의 파일 안에 여러 오브젝트와 재료 정보를 포함할 수 있어 편리하다.')
doc.add_paragraph('PrusaSlicer가 이미 STL, OBJ, 3MF, AMF, STEP을 지원하므로 기존 로더를 그대로 활용한다.')

doc.add_heading('2.5 설계 원칙', level=2)
for b in [
    '슬라이서는 단순하고 안정적인 구조를 유지',
    'Boolean 연산 등 복잡한 메시 연산은 피한다 (외부 CAD가 담당)',
    'STL 모델을 수정하지 않고 그대로 슬라이스',
    '기존 기능 위에 추가하는 방식 (처음부터 새로 만들지 않는다)',
    '레이어 시퀀스는 슬라이서/vgui/펌웨어 3곳 모두 동일하게',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ============================================================
# 3. 하드웨어 제어 파라미터
# ============================================================
doc.add_heading('3. 하드웨어 제어 파라미터', level=1)

doc.add_heading('3.1 듀얼 노즐 오프셋', level=2)
doc.add_paragraph(
    '두 개의 레진 노즐은 물리적으로 서로 다른 위치에 있다. '
    '슬라이서는 Tool 0과 Tool 1에 대해 노즐 오프셋 보정을 적용해야 한다. '
    '이는 슬라이서(GCode 좌표 보정)와 프린터 캘리브레이션 양쪽 모두에 영향을 준다.')
doc.add_paragraph()
add_table(doc, [
    ('파라미터', '타입', '단위', '설명'),
    ('nozzle_offset_x', 'float', 'mm', 'T1 노즐의 X 오프셋 (T0 기준)'),
    ('nozzle_offset_y', 'float', 'mm', 'T1 노즐의 Y 오프셋 (T0 기준)'),
    ('nozzle_offset_z', 'float', 'mm', 'T1 노즐의 Z 오프셋 (T0 기준, 선택)'),
])

doc.add_heading('3.2 디스펜서 제어 (공기압)', level=2)
doc.add_paragraph(
    '공기압 기반 디스펜서는 FDM의 스테퍼 모터 E축과 완전히 다른 제어 방식이다. '
    '토출 시작과 종료 시 지연이 발생하며, 잔류 압력에 의한 추가 토출이 있을 수 있다. '
    '이를 보정하기 위한 디스펜서 제어 파라미터가 필요하다.')
doc.add_paragraph()
add_table(doc, [
    ('파라미터', '타입', '단위', '설명'),
    ('resin_dispense_start_delay', 'int', 'ms', '토출 시작 지연 시간'),
    ('resin_dispense_stop_delay', 'int', 'ms', '토출 종료 지연 시간'),
    ('resin_pressure', 'int', 'kPa', '디스펜서 공기압'),
    ('resin_nozzle_diameter', 'float', 'mm', '도포 노즐 직경'),
    ('resin_flow_rate', 'float', 'mm3/s', '레진 토출량 (선택)'),
])

doc.add_heading('3.3 블레이드 제어', level=2)
doc.add_paragraph(
    '레진 도포 후 블레이드가 레이어를 평탄화한다. '
    '블레이드 동작 방식은 편도(one-way) 또는 왕복(round-trip)으로 설정할 수 있다. '
    '이 프린터의 핵심 메커니즘이므로 세밀한 제어가 필요하다.')
doc.add_paragraph()
add_table(doc, [
    ('파라미터', '타입', '단위', '설명'),
    ('blade_speed', 'float', 'mm/s', '블레이드 이동 속도'),
    ('blade_direction', 'enum', '-', 'one_way / round_trip'),
    ('blade_pass_count', 'int', '회', '블레이드 왕복/편도 횟수'),
    ('blade_height', 'float', 'mm', '블레이드 높이 (레진 표면 기준)'),
    ('blade_start_position', 'float', 'mm', '블레이드 시작 X 위치'),
    ('blade_end_position', 'float', 'mm', '블레이드 종료 X 위치'),
])

doc.add_heading('3.4 LED/프로젝터 제어 + Flat Field 보정', level=2)
doc.add_paragraph('프로젝터 관련 파라미터:')
add_table(doc, [
    ('파라미터', '타입', '단위', '설명'),
    ('exposure_time', 'float', '초', '일반 레이어 노출 시간'),
    ('first_layer_exposure_time', 'float', '초', '첫 레이어 노출 시간 (접착력 강화)'),
    ('resolution_x', 'int', 'px', '프로젝터 가로 해상도'),
    ('resolution_y', 'int', 'px', '프로젝터 세로 해상도'),
    ('pixel_size', 'float', 'um', '픽셀 물리적 크기'),
])
doc.add_paragraph()
doc.add_paragraph(
    'Flat Field 보정: 프로젝터 밝기가 가장자리로 갈수록 어두워지는 불균일 문제를 보정한다. '
    'flat_field_mask.png를 출력 ZIP에 포함하여 vgui에서 마스크 이미지에 적용할 수 있다. '
    '초기 버전에서는 선택 기능(후순위)이다.')

doc.add_page_break()

# ============================================================
# 4. PrusaSlicer 활용 전략
# ============================================================
doc.add_heading('4. PrusaSlicer 활용 전략', level=1)

doc.add_heading('4.1 SLA에서 가져올 것 (몸통)', level=2)
add_table(doc, [
    ('기능', '소스 위치', '용도'),
    ('3D 뷰어 + 모델 배치', 'src/slic3r/GUI/', '기존 UI 그대로'),
    ('Z 슬라이싱 -> ExPolygon', 'PrintObjectSlice.cpp', '레이어 단면 생성'),
    ('ExPolygon -> PNG 래스터화', 'SLA/RasterBase.*', '마스크 이미지 생성'),
    ('SLA 서포트', 'src/libslic3r/SLA/', '서포트 (레진 방식)'),
    ('SL1 ZIP 출력', 'Format/SL1.*', 'ZIP 패키징 참고'),
    ('레이어 높이/노출 설정', 'PrintConfig.hpp', '경화 설정'),
    ('Zipper 클래스', 'Zipper.*', 'ZIP 생성'),
])

doc.add_heading('4.2 FDM에서 가져올 것 (추가)', level=2)
add_table(doc, [
    ('기능', '소스 위치', '용도'),
    ('멀티 익스트루더 -> 멀티 Tool', 'PrintConfig, PrintObject', 'Object별 Tool 지정'),
    ('인필 패턴 -> 도포 경로', 'src/libslic3r/Fill/', '도포 패턴 활용'),
    ('GCode 생성', 'src/libslic3r/GCode/*', 'Klipper용 도포 GCode'),
    ('GCodeWriter (gcfKlipper)', 'GCodeWriter.*', 'Klipper 명령어'),
])

doc.add_heading('4.3 제거할 것 (FDM/Prusa 전용)', level=2)
doc.add_paragraph('UI에서 제거하거나 숨겨야 하는 FDM 전용 기능:')
add_table(doc, [
    ('제거 대상', '관련 파일/UI', '이유'),
    ('Prusa 계정 로그인', 'UserAccount*, LoginDialog*', 'Prusa 전용'),
    ('PrusaConnect', 'PrusaConnect.cpp', 'Prusa 전용'),
    ('Configuration Sources', 'ConfigWizard.cpp', 'Prusa 전용'),
    ('프리셋 자동 업데이트', 'PresetUpdater*', 'Prusa 전용'),
    ('기존 프린터 프로파일', 'resources/profiles/', '불필요'),
    ('온도 제어 (Temperature)', 'UI + GCode', '레진에 불필요'),
    ('필라멘트 설정 (Filament)', 'UI + PrintConfig', '레진 사용'),
    ('냉각 팬 (Cooling Fan)', 'CoolingBuffer + UI', '레진에 불필요'),
    ('Extrusion Multiplier', 'UI + PrintConfig', 'E축 미사용 (공기압)'),
    ('리트랙션 (Retraction)', 'UI + GCode', 'E축 미사용'),
    ('벽 정밀 생성 (Perimeter)', 'PerimeterGenerator, Arachne', '도포는 정밀 벽 불필요'),
    ('아이로닝 (Ironing)', 'UI + GCode', '블레이드가 대신함'),
    ('와이프 타워 (Wipe Tower)', 'UI + GCode', '멀티 레진 방식 다름'),
])

doc.add_heading('4.4 참고 오픈소스', level=2)
add_table(doc, [
    ('프로젝트', '역할', '대응', 'GitHub'),
    ('LumenSlicer', 'PrusaSlicer 포크 레진 슬라이서', '커스텀 슬라이서', 'Open-Resin-Alliance/LumenSlicer'),
    ('Odyssey', 'SL1 처리 + Klipper 통신', 'vgui', 'Open-Resin-Alliance/Odyssey'),
    ('Orion', 'Flutter 프린터 GUI', 'vgui', 'Open-Resin-Alliance/Orion'),
])

doc.add_page_break()

# ============================================================
# 5. 개발 Phase
# ============================================================
doc.add_heading('5. 개발 Phase', level=1)

# Phase -1
doc.add_heading('Phase -1: PrusaSlicer 사용법 학습', level=2)
bold_then(doc, '목표: ', '개발 전에 PrusaSlicer를 사용자로서 충분히 사용해본다.')
doc.add_paragraph('슬라이서 개발자는 슬라이서를 먼저 잘 써봐야 한다:')
for b in [
    'FDM 모드: STL 로드, 슬라이싱, GCode 미리보기, 멀티 익스트루더 설정',
    'SLA 모드: 레진 슬라이싱, 레이어 이미지 확인, 서포트 생성',
    '프리셋 시스템: 프린터/재료/프로파일 설정 구조 이해',
    '3MF 파일 로드, 멀티 오브젝트 배치',
]:
    doc.add_paragraph(b, style='List Bullet')

# Phase 0
doc.add_heading('Phase 0: 빌드 환경 구축', level=2)
bold_then(doc, '목표: ', 'PrusaSlicer를 그대로 빌드하고 실행할 수 있는 환경.')
doc.add_paragraph('C++17 (MSVC 2019+), CMake 3.13+, Git. deps/ 자동 빌드.')
p = doc.add_paragraph()
r = p.add_run('빌드가 안 되면 아무것도 못 한다.')
r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

# Phase 1
doc.add_heading('Phase 1: 불필요 기능 제거 + 브랜딩', level=2)
bold_then(doc, '목표: ', 'Prusa 전용 + FDM 전용 기능 비활성화. 섹션 4.3 참고.')

# Phase 2
doc.add_heading('Phase 2: 멀티 레진 Tool 매핑', level=2)
bold_then(doc, '목표: ', 'Object -> Material -> Tool -> Nozzle 매핑 구현.')
doc.add_paragraph('PrusaSlicer 멀티 익스트루더 기능을 리네이밍하여 활용.')

# Phase 3
doc.add_heading('Phase 3: 마스크 이미지(PNG) 생성', level=2)
bold_then(doc, '목표: ', 'SLA RasterBase 참고. ExPolygon -> PNG. 전체 합쳐서 1장.')

# Phase 4
doc.add_heading('Phase 4: 도포 경로 GCode 생성', level=2)
bold_then(doc, '목표: ', 'Tool별 도포 영역에 대한 GCode 생성.')
doc.add_paragraph('포함 사항: 듀얼 노즐 오프셋 보정, 디스펜서 시작/정지 지연 보정, 블레이드 GCode')
doc.add_paragraph()
doc.add_paragraph('레이어별 GCode 예시:')
add_code(doc, [
    '; Layer 42',
    'T0                              ; 모델레진 노즐',
    'M400                            ; 디스펜서 시작 대기',
    'G1 X10.0 Y15.0 F1200            ; T0 도포 경로',
    'G1 X50.0 Y15.0',
    'G1 X50.0 Y45.0',
    'M400                            ; 디스펜서 정지 대기',
    '',
    'T1                              ; 템프레진 노즐 (오프셋 보정 적용)',
    'M400',
    'G1 X25.0 Y20.0 F1200            ; T1 도포 경로',
    'G1 X40.0 Y35.0',
    'M400',
    '',
    '; Blade sweep',
    'G1 X0 F600                      ; 블레이드 시작 위치',
    'G1 X200 F{blade_speed}          ; 블레이드 스윕',
    '',
    '; Expose',
    '; DISPLAY_IMAGE layers/0042.png (vgui가 GCode 파싱하여 처리)',
    'G4 P2500                        ; 노출 대기 2.5초',
    '',
    '; Z lift',
    'G1 Z2.15 F300                   ; 다음 레이어',
])

# Phase 5
doc.add_heading('Phase 5: 출력 포맷 패키징 (ZIP)', level=2)
bold_then(doc, '목표: ', 'GCode + PNG + manifest.json을 ZIP으로 패키징. SL1Archive 참고.')

# Phase 6
doc.add_heading('Phase 6: vgui 연동 + 실 프린트 테스트', level=2)
bold_then(doc, '목표: ', 'vgui가 새 포맷을 읽고 프린팅 실행. USB 테스트.')

doc.add_page_break()

# ============================================================
# 6. 수정 파일 + 코드 탐색 가이드
# ============================================================
doc.add_heading('6. 주요 수정 파일 + 코드 탐색 가이드', level=1)

doc.add_heading('초기 코드 분석 대상 (반드시 먼저 읽을 파일)', level=2)
add_table(doc, [
    ('파일', '역할', '왜 읽어야 하는가'),
    ('src/libslic3r/Print.hpp,cpp', '프린트 프로세스 관리', '전체 파이프라인 이해'),
    ('src/libslic3r/PrintObject.cpp', '객체별 슬라이싱', '슬라이싱 단계 흐름'),
    ('src/libslic3r/PrintObjectSlice.cpp', '메시->레이어 변환', '겹침 처리 로직'),
    ('src/libslic3r/GCode.hpp,cpp', 'GCode 생성 엔진', 'GCode 출력 방식'),
    ('src/libslic3r/SLA/RasterBase.hpp,cpp', 'ExPolygon->래스터', '마스크 이미지 생성 방법'),
    ('src/libslic3r/Format/SL1.hpp,cpp', 'SL1 ZIP 출력', '출력 포맷 구조'),
    ('src/libslic3r/PrintConfig.hpp', '300+ 설정', '파라미터 추가 방법'),
    ('src/libslic3r/Config.hpp', 'PrinterTechnology enum', 'ptHybrid 추가 위치'),
])

doc.add_paragraph()
doc.add_heading('전체 수정 파일 목록', level=2)
add_table(doc, [
    ('파일', '수정 내용', '우선순위'),
    ('Config.hpp', 'PrinterTechnology에 ptHybrid 추가', '★★★'),
    ('PrintConfig.hpp', '하이브리드 파라미터 (노즐/디스펜서/블레이드/LED)', '★★★'),
    ('SLA/RasterBase.*', '마스크 이미지 생성 참고/활용', '★★★'),
    ('Format/SL1.*', 'HybridArchiveWriter 작성 참고', '★★★'),
    ('GCode.*, GCodeWriter.*', '도포 GCode + 블레이드 + 노즐 오프셋', '★★☆'),
    ('Fill/', '인필 패턴 -> 도포 경로', '★★☆'),
    ('Print.*, PrintObject.*', '하이브리드 파이프라인', '★★☆'),
    ('Preset.*', '프리셋 하이브리드 타입', '★★☆'),
    ('resources/profiles/', '프린터 프로파일', '★★☆'),
    ('GUI/MainFrame.*', '브랜딩', '★☆☆'),
    ('GUI/ConfigWizard.*', '하이브리드 선택', '★☆☆'),
    ('GUI/UserAccount*', '로그인 비활성화', '★☆☆'),
])

doc.add_page_break()

# ============================================================
# 7. vgui 수정 사항
# ============================================================
doc.add_heading('7. vgui (프린터 측) 수정 사항', level=1)
doc.add_paragraph(
    '슬라이서 출력 포맷 변경에 따라 vgui 수정 필요. '
    '핵심 로직(Moonraker 통신, 프로젝터 표시)은 유지.')
doc.add_paragraph()
add_table(doc, [
    ('바뀌는 것', '안 바뀌는 것'),
    ('ZIP 파일 파싱', 'Moonraker API 호출'),
    ('manifest.json 읽기', '모터 제어 흐름'),
    ('PNG 경로 규칙 (layers/*.png)', '프로젝터 이미지 표시'),
    ('GCode 내 커스텀 명령어 처리', 'USB 감지/마운트'),
    ('Flat Field 보정 적용 (선택)', 'UI 레이아웃'),
])
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('출력 파일 포맷(ZIP 구조, manifest 스펙)을 먼저 확정한 후 슬라이서와 vgui를 동시 개발.')
r.bold = True

doc.add_page_break()

# ============================================================
# 8. 핵심 원칙 + 일정
# ============================================================
doc.add_heading('8. 핵심 원칙 + 일정', level=1)

doc.add_heading('핵심 원칙', level=2)
principles = [
    ('레진 슬라이서가 몸통', '마스크 이미지(PNG)가 형상 결정. 도포 GCode는 보조.'),
    ('기존 위에 추가', 'PrusaSlicer 포크 후 점진적 수정. 새로 만들지 않는다.'),
    ('STL을 수정하지 않는다', '외부 CAD에서 좌표 확정. Boolean 연산 회피.'),
    ('출력 포맷 먼저 확정', '슬라이서(PC)와 vgui(프린터)가 같은 포맷을 써야 의미 있다.'),
    ('레이어 시퀀스 통일', '도포->블레이드->노광->경화->Z이동. 3곳 모두 동일.'),
    ('이상적 시나리오로 개발', '물리적 문제(혼합, 쏠림)는 하드웨어/재료가 해결.'),
    ('한 번에 하나씩', '각 Phase 완료 + 테스트 후 다음.'),
]
for i, (t, d) in enumerate(principles, 1):
    p = doc.add_paragraph()
    r = p.add_run(f'{i}. {t}')
    r.bold = True
    r.font.size = Pt(11)
    doc.add_paragraph(d)

doc.add_heading('권장 일정', level=2)
add_table(doc, [
    ('기간', 'Phase', '주요 작업'),
    ('0주차', 'Phase -1', 'PrusaSlicer 사용법 학습 (FDM/SLA 모드, 멀티 익스트루더)'),
    ('1주차', 'Phase 0', '빌드 환경 구축, 정상 빌드/실행 확인'),
    ('2주차', 'Phase 1', '불필요 기능 비활성화, 브랜딩, FDM 전용 UI 제거'),
    ('3-4주차', 'Phase 2', '멀티 레진 Tool 매핑 (Object->Material->Tool->Nozzle)'),
    ('5-6주차', 'Phase 3', '마스크 이미지 PNG 생성 (RasterBase)'),
    ('7-8주차', 'Phase 4', '도포 GCode + 노즐 오프셋 + 디스펜서 + 블레이드'),
    ('9-10주차', 'Phase 5', 'ZIP 패키징 + manifest.json (SL1Archive)'),
    ('11-12주차', 'Phase 6', 'vgui 연동 + 실 프린트 테스트'),
])

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('-- End --')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

output_path = r'c:\Users\JoWooHyun\Documents\PrusaSlicer\커스텀_슬라이서_개발계획서_v3.docx'
doc.save(output_path)
print(f'Done: {output_path}')
